import os
import io
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from docx import Document
from typing import List, Dict


def extract_structured_chunks(file_bytes: bytes, filename: str, enable_ocr: bool = True) -> List[Dict]:
    ext = os.path.splitext(filename)[1].lower()
    print(f"[DEBUG] Estensione file: {ext}", flush=True)

    if ext == ".pdf":
        return extract_from_pdf(file_bytes, enable_ocr)
    elif ext == ".docx":
        return extract_from_docx(file_bytes)
    elif ext == ".txt":
        return extract_from_txt(file_bytes)
    else:
        raise ValueError(f"❌ Estensione non supportata: {ext}")


def extract_from_pdf(file_bytes: bytes, enable_ocr: bool = True) -> List[Dict]:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    chunks = []

    print(f"[DEBUG] Pagine PDF: {len(doc)}", )

    for page_index, page in enumerate(doc):
        print(f"[DEBUG] → Pagina {page_index + 1}")

        # === Testo ===
        blocks = page.get_text("dict")["blocks"]
        print(f"[DEBUG]   Blocchi totali nella pagina: {len(blocks)}")

        for block in blocks:
            if block["type"] == 0:
                paragraph = ""
                max_font_size = 0
                for line in block["lines"]:
                    for span in line["spans"]:
                        paragraph += span["text"] + " "
                        max_font_size = max(max_font_size, span.get("size", 0))

                if paragraph.strip():
                    chunk_type = "heading" if max_font_size > 15 else "paragraph"
                    print(f"[DEBUG]   → Aggiunto blocco: {chunk_type} | {paragraph.strip()[:50]}...")
                    chunks.append({
                        "type": chunk_type,
                        "text": paragraph.strip(),
                        "level": 1 if chunk_type == "heading" else None
                    })

        # === Immagini + OCR ===
        if enable_ocr:
            image_list = page.get_images(full=True)
            print(f"[DEBUG]   Immagini trovate nella pagina: {len(image_list)}")

            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image = Image.open(io.BytesIO(image_bytes))
                    ocr_text = pytesseract.image_to_string(image)
                    print(f"[DEBUG]   → OCR immagine {img_index + 1}: {ocr_text[:50]}...")

                    chunks.append({
                        "type": "image",
                        "text": "🖼️ Immagine nel documento",
                        "caption": None,
                        "ocr_text": ocr_text.strip() or None
                    })
                except Exception as e:
                    print(f"[ERROR]   ❌ OCR fallito per immagine {img_index + 1}: {e}")
                    chunks.append({
                        "type": "image",
                        "text": "🖼️ Immagine nel documento",
                        "caption": None,
                        "ocr_text": None
                    })

    print(f"[DEBUG] → Chunks totali PDF: {len(chunks)}")
    return chunks


def extract_from_docx(file_bytes: bytes) -> List[Dict]:
    doc = Document(io.BytesIO(file_bytes))
    chunks = []

    print(f"[DEBUG] Paragrafi nel DOCX: {len(doc.paragraphs)}")

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style = para.style.name.lower()
        if "heading" in style:
            try:
                level = int(style.replace("heading", "").strip())
            except:
                level = 1
            print(f"[DEBUG] → Heading livello {level}: {text[:50]}")
            chunks.append({"type": "heading", "text": text, "level": level})
        else:
            print(f"[DEBUG] → Paragrafo: {text[:50]}")
            chunks.append({"type": "paragraph", "text": text, "level": None})

    print(f"[DEBUG] → Chunks totali DOCX: {len(chunks)}")
    return chunks


def extract_from_txt(file_bytes: bytes) -> List[Dict]:
    text = file_bytes.decode("utf-8", errors="ignore")
    lines = text.splitlines()
    chunks = []

    print(f"[DEBUG] Linee nel TXT: {len(lines)}")

    for line in lines:
        if line.strip():
            print(f"[DEBUG] → Riga: {line.strip()[:50]}")
            chunks.append({"type": "paragraph", "text": line.strip(), "level": None})

    print(f"[DEBUG] → Chunks totali TXT: {len(chunks)}")
    return chunks
